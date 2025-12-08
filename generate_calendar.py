"""
Calendar Generator Script

This script generates a writable calendar for a specified year in Microsoft Word format (.docx).
It creates a landscape-oriented document with one page per month, featuring a table-based calendar layout
where users can write notes in each day's cell. Weekends are highlighted in orange for easy identification.
US federal holidays, Christian holidays (Easter, Good Friday, Pentecost, Mother's Day, Father's Day), LDS Church-specific holidays
(General Conference on the first weekend in April and October, Pioneer Day), California school holidays
(Cesar Chavez Day), and key LDS Church history events are marked with their names in smaller italic text below the date.
If multiple events occur on the same date, they are separated by "new line".

Usage:
    python generate_calendar.py [year]

    - year: Optional integer argument for the year (e.g., 2025). Defaults to the current year if not provided.

Dependencies:
    - python-docx: For creating and manipulating Word documents.
    - calendar: Standard library module for calendar-related functions.

Output:
    A .docx file named 'Calendar_{year}_Writable.docx' in the current directory.

Author: Edward L. Thomas
Date: December 7, 2025
Update: December 7, 2025
Version: 2.0
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: The 'python-docx' library is required but not installed.")
    print("Please install it using: pip install python-docx")
    sys.exit(1)

import calendar
import sys
import datetime
from datetime import date, timedelta

# Function to calculate the date of Easter for a given year using the Meeus/Jones/Butcher algorithm
def calculate_easter(year):
    a = year % 19  # Golden number (position in 19-year Metonic cycle)
    b = year // 100  # Century
    c = year % 100  # Year within the century
    d = b // 4  # Number of leap centuries
    e = b % 4  # Remainder of century divided by 4
    f = (b + 8) // 25  # Solar cycle adjustment
    g = (b - f + 1) // 3  # Century correction factor
    h = (19 * a + b - d - g + 15) % 30  # Epact (age of the Moon on Jan 1)
    i = c // 4  # Number of leap years in the century
    k = c % 4  # Remainder of year within century divided by 4
    l = (32 + 2 * e + 2 * i - h - k) % 7  # Day of week for Paschal full moon
    m = (a + 11 * h + 22 * l) // 451  # Correction factor for date bounds
    month = (h + l - 7 * m + 114) // 31  # Month of Easter (3=March, 4=April)
    day = ((h + l - 7 * m + 114) % 31) + 1  # Day of the month
    return date(year, month, day)

# Function to get the first Sunday of a given month and year
def first_sunday(year, month):
    first_day = date(year, month, 1)
    weekday = first_day.weekday()  # 0=Monday, 6=Sunday
    days_to_sunday = (6 - weekday) % 7
    return first_day + timedelta(days=days_to_sunday)

# Function to get the holiday names for a given date, if any
def get_holiday_name(d):
    holidays = []
    
    # US Federal Holidays
    if d.month == 1 and d.day == 1:
        holidays.append("New Year's Day")
    if d.month == 1 and d.weekday() == 0:  # Monday
        # Third Monday in January
        jan1 = date(d.year, 1, 1)
        first_mon = jan1 + timedelta(days=(7 - jan1.weekday()) % 7)
        third_mon = first_mon + timedelta(days=14)
        if d == third_mon:
            holidays.append("Martin Luther King Jr. Day")
    if d.month == 2 and d.weekday() == 0:  # Monday
        # Third Monday in February
        feb1 = date(d.year, 2, 1)
        first_mon = feb1 + timedelta(days=(7 - feb1.weekday()) % 7)
        third_mon = first_mon + timedelta(days=14)
        if d == third_mon:
            holidays.append("Presidents' Day")
    if d.month == 5 and d.weekday() == 0:  # Monday
        # Last Monday in May
        may31 = date(d.year, 5, 31)
        last_mon = may31 - timedelta(days=may31.weekday())
        if d == last_mon:
            holidays.append("Memorial Day")
    if d.month == 7 and d.day == 4:
        holidays.append("Independence Day")
    if d.month == 9 and d.weekday() == 0:  # Monday
        # First Monday in September
        sep1 = date(d.year, 9, 1)
        first_mon = sep1 + timedelta(days=(7 - sep1.weekday()) % 7)
        if d == first_mon:
            holidays.append("Labor Day")
    if d.month == 10 and d.weekday() == 0:  # Monday
        # Second Monday in October
        oct1 = date(d.year, 10, 1)
        first_mon = oct1 + timedelta(days=(7 - oct1.weekday()) % 7)
        second_mon = first_mon + timedelta(days=7)
        if d == second_mon:
            holidays.append("Columbus Day")
    if d.month == 11 and d.day == 11:
        holidays.append("Veterans Day")
    if d.month == 11 and d.weekday() == 3:  # Thursday
        # Fourth Thursday in November
        nov1 = date(d.year, 11, 1)
        first_thu = nov1 + timedelta(days=(3 - nov1.weekday()) % 7)
        fourth_thu = first_thu + timedelta(days=21)
        if d == fourth_thu:
            holidays.append("Thanksgiving")
    if d.month == 12 and d.day == 25:
        holidays.append("Christmas")
    if d.month == 3 and d.day == 31:
        holidays.append("Cesar Chavez Day")

    # Christian Holidays
    easter = calculate_easter(d.year)
    if d == easter - timedelta(days=2):
        holidays.append("Good Friday")
    if d == easter:
        holidays.append("Easter")
    if d == easter + timedelta(days=49):
        holidays.append("Pentecost")
    # Mother's Day: Second Sunday in May
    if d.month == 5 and d.weekday() == 6:  # Sunday
        may1 = date(d.year, 5, 1)
        first_sun = may1 + timedelta(days=(6 - may1.weekday()) % 7)
        second_sun = first_sun + timedelta(days=7)
        if d == second_sun:
            holidays.append("Mother's Day")
    # Father's Day: Third Sunday in June
    if d.month == 6 and d.weekday() == 6:  # Sunday
        jun1 = date(d.year, 6, 1)
        first_sun_jun = jun1 + timedelta(days=(6 - jun1.weekday()) % 7)
        third_sun_jun = first_sun_jun + timedelta(days=14)
        if d == third_sun_jun:
            holidays.append("Father's Day")

    # LDS Church Holidays
    april_sunday = first_sunday(d.year, 4)
    oct_sunday = first_sunday(d.year, 10)
    if d == april_sunday - timedelta(days=1) or d == april_sunday:
        holidays.append("General Conference")
    if d == oct_sunday - timedelta(days=1) or d == oct_sunday:
        holidays.append("General Conference")
    if d.month == 7 and d.day == 24:
        holidays.append("Pioneer Day")

    # LDS Church History Events
    if d.month == 9 and d.day == 21:
        holidays.append("First Vision")
    if d.month == 4 and d.day == 6:
        holidays.append("Church Organization")
    if d.month == 3 and d.day == 27:
        holidays.append("Kirtland Temple Dedication")
    if d.month == 4 and d.day == 3:
        holidays.append("First Presidency Organized")
    if d.month == 6 and d.day == 27:
        holidays.append("Joseph Smith Martyrdom")

    return holidays

# Parse command line arguments to determine the year for the calendar
if len(sys.argv) > 1:
    try:
        year = int(sys.argv[1])  # Attempt to convert the argument to an integer
    except ValueError:
        print("Error: Year must be a valid integer")  # Inform user of invalid input
        sys.exit(1)  # Exit the program with error code
else:
    year = datetime.datetime.now().year  # Default to the current year if no argument is provided

# Initialize a new Word document
doc = Document()
doc.core_properties.title = f"{year} Calendar - Writable Version"  # Set the document title

# Configure the calendar module to start weeks on Sunday (default is Monday)
calendar.setfirstweekday(calendar.SUNDAY)

# Configure the document's page layout for landscape orientation
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE  # Change to landscape mode
section.page_width = Inches(11)  # Standard letter width in landscape
section.page_height = Inches(8.5)  # Standard letter height in landscape
section.top_margin = Inches(0.1)  # Minimal top margin
section.bottom_margin = Inches(0.1)  # Minimal bottom margin
section.left_margin = Inches(0.5)  # Left margin
section.right_margin = Inches(0.5)  # Right margin

# Loop through each month (1 to 12) to create a page for each
for month in range(1, 13):
    # Add a centered heading for the month and year
    heading = doc.add_heading(level=0)
    run = heading.add_run(f"{calendar.month_name[month]} {year}")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create a 7x7 table: 1 row for weekdays, up to 6 rows for weeks
    table = doc.add_table(rows=7, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Center the table on the page
    table.allow_autofit = False  # Disable auto-fitting to maintain custom widths

    # Set uniform column widths suitable for landscape orientation
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(1.4)  # Width adjusted for landscape layout

    # Define a function to remove all borders from the table for a clean look
    def remove_table_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr.append(tblBorders)

    remove_table_borders(table)  # Apply the border removal to the current table

    # Define a function to add thin gray borders to individual cells
    def add_cell_borders(cell):
        tcPr = cell._element.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            cell._element.insert(0, tcPr)
        
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')   # 4 = 0.5 pt
            border.set(qn('w:color'), 'D3D3D3')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    # Populate the header row with weekday abbreviations
    headers = ["SUN", "MON", "TUE", "WED", "THU", "FRI", "SAT"]
    hdr_cells = table.rows[0].cells
    for i, day_name in enumerate(headers):
        cell = hdr_cells[i]
        p = cell.paragraphs[0]
        p.text = day_name
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the text
        run = p.runs[0]
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(80, 80, 80)  # Gray color for headers
        # Add borders to header cells for definition
        add_cell_borders(cell)

    # Get the calendar data for the current month as a list of weeks
    weeks = calendar.monthcalendar(year, month)

    for week_idx in range(6):  # Up to 6 weeks in a month
        row_cells = table.rows[week_idx + 1].cells
        # Set row height to provide space for writing notes
        table.rows[week_idx + 1].height = Inches(1.0)

        for day_idx in range(7):  # 7 days in a week
            cell = row_cells[day_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left-align text for notes

            # Add top padding inside the cell for better spacing
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)

            if week_idx < len(weeks) and weeks[week_idx][day_idx] != 0:
                day = weeks[week_idx][day_idx]
                date_obj = date(year, month, day)
                holidays = get_holiday_name(date_obj)
                run = p.add_run(str(day))  # Add the day number
                run.font.size = Pt(18)
                run.bold = True

                if holidays:
                    holiday_text = "\n".join(holidays)
                    run.add_break()  # Add a line break
                    run2 = p.add_run(holiday_text)  # Add the holiday names separated by new lines
                    run2.font.size = Pt(10)
                    run2.italic = True

                # Highlight weekends in soft orange for visual distinction
                if day_idx in (0, 6):  # Sunday (0) or Saturday (6)
                    run.font.color.rgb = RGBColor(230, 138, 0)  # Warm orange
                
                # Add borders only to cells containing dates
                add_cell_borders(cell)
            elif week_idx == 0:
                # Add borders to blank cells at the start of the month (first week only)
                add_cell_borders(cell)
            # Other empty cells remain without borders for a clean look

    # Add a page break after each month except the last one
    if month < 12:
        doc.add_page_break()

# Save the completed document to a file
output_path = f"Calendar_{year}_Writable.docx"
doc.save(output_path)

print(f"Calendar successfully saved as: {output_path}")  # Confirm successful save


